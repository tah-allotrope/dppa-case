"""
Build lessons/DPPA_Worksheets_and_Answers.docx
==============================================

A bilingual (EN | VI) Word document that pairs a blank 5-line compute
worksheet with its worked totals for each workshop scenario (S1 matched,
S2 shortfall, S3 excess), followed by a negotiation block and a 3-case
comparison summary table.

The visual shell (A4, 0.75in margins, blue banner / subhead / callout /
footer styling) is inherited by copying the reference template
``lessons/DPPA_Scenario_Answer_Summary.docx`` and clearing its scenario
body. The numeric totals are read from the engine-generated spine packs
(``assets/teaching/spine-s{1,2,3}.json`` — see PHASE-04 of
plans/2026-07-17-prose-parity-second-pipeline-plan.md), not hand-typed; the
module refuses to build from drifted spines (see the assertions right
after ``load_spines``).

Run:  python build_worksheet_answer_docx.py
"""
from __future__ import annotations

import json
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ----------------------------------------------------------------------
# Paths
# ----------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "lessons" / "DPPA_Scenario_Answer_Summary.docx"
OUTPUT = ROOT / "lessons" / "DPPA_Worksheets_and_Answers.docx"

# ----------------------------------------------------------------------
# Style constants — observed from the reference (PHASE-01)
# ----------------------------------------------------------------------
BANNER_FILL = "1F4E79"        # title / scenario banner background
SUBTITLE_FILL = "D6E4F0"      # banner subtitle + callout fill
HEADER_FILL = "1F4E79"        # data-table header row
NOTE_FILL = "FFFFFF"          # white callout/note fill
SECTION_COLOR = "1F4E79"      # section heading color
SUBHEAD_COLOR = "2E75B6"      # sub-heading color
FOOTER_COLOR = "888888"       # footer text
EN_HEADER_FILL = "1F4E79"     # used in data-table header rows

# Word font sizes are stored in half-points (sz=36 -> 18pt)
SZ_TITLE = "36"               # 18pt — banner title
SZ_SUBTITLE = "28"            # 14pt — banner subtitle
SZ_LEGAL = "18"               # 9pt  — banner legal basis
SZ_SECTION = "24"             # 12pt — section heading
SZ_BODY = "22"                # 11pt — body / table cells
SZ_FOOTER = "16"              # 8pt  — footer

# Page geometry: 9360 dxa ~ 6.5in usable width on A4 with 0.75in margins
TABLE_WIDTH = "9360"
COL_LABEL = "1620"            # ~1.13in
COL_CALC = "5400"             # ~3.75in (Calculation + EN|VI pair in one cell)
COL_ANSWER = "2340"           # ~1.63in (VND / month fillable)

# Footer line (verbatim from the reference, bilingual)
FOOTER_TEXT = (
    "Prepared for CEBA 2026 Training  |  "
    "Được chuẩn bị cho Chương trình Đào tạo CEBA 2026  |  "
    "Allotrope Partners Vietnam  |  "
    "Based on Decree 57/2025/NĐ-CP & Decree 243/2026/NĐ-CP"
)

# ----------------------------------------------------------------------
# Spine loading — engine-generated numbers (PHASE-04)
# ----------------------------------------------------------------------
def load_spines(root: Path) -> dict[str, dict]:
    """Load assets/teaching/spine-s{1,2,3}.json, keyed 's1'/'s2'/'s3'."""
    spines: dict[str, dict] = {}
    for key in ("s1", "s2", "s3"):
        path = root / "assets" / "teaching" / f"spine-{key}.json"
        if not path.exists():
            raise FileNotFoundError(f"Missing spine export: {path}")
        spines[key] = json.loads(path.read_text(encoding="utf-8"))
    return spines


def fmt_vnd(value: int) -> str:
    """Comma-grouped absolute value, U+2212 minus prefix when negative."""
    return f"−{abs(value):,}" if value < 0 else f"{value:,}"


def fmt_signed_vnd(value: int) -> str:
    """Like fmt_vnd but with an explicit '+' for positives (CfD lines)."""
    return f"−{abs(value):,}" if value < 0 else f"+{value:,}"


def fmt_effective(c_kh_vnd: int, total_kwh: int) -> str:
    return f"≈ {round(c_kh_vnd / total_kwh):,} VND/kWh"


def fmt_int(value: int) -> str:
    return f"{value:,}"


SPINES = load_spines(ROOT)

# Refuse to build from drifted spines — same anchors as
# app/scripts/export-spine.mjs and Specification S1 of the plan.
assert SPINES["s1"]["bill"]["cKh"]["vnd"] == 9_063_196_000
assert SPINES["s2"]["bill"]["cEvn"]["vnd"] == 19_628_262_400
assert SPINES["s2"]["bill"]["cKh"]["vnd"] == 18_828_262_400
assert SPINES["s3"]["bill"]["cEvn"]["vnd"] == 8_304_644_000
assert SPINES["s3"]["bill"]["cKh"]["vnd"] == 9_054_644_000
assert SPINES["s3"]["bill"]["lines"]["cfd"]["vnd"] == 750_000_000

# ----------------------------------------------------------------------
# Canonical (small) constants — shared across scenarios, sourced from S1's
# spine inputs since fees/loss-factor are identical for all three scenarios.
# ----------------------------------------------------------------------
K = 1.026
K_PP = SPINES["s1"]["inputs"]["lossFactorKppOnly"]
K_KPP = SPINES["s1"]["inputs"]["lossFactorPrecise"]   # k * K_pp (precise)
SERVICE = SPINES["s1"]["inputs"]["serviceFee"]
CLEARING = SPINES["s1"]["inputs"]["clearingFee"]
FEES = SERVICE + CLEARING                             # service + clearing
RETAIL = SPINES["s1"]["inputs"]["retailTariff"]


def build_scenario(key: str, en_title: str, vi_title: str) -> dict:
    """Build one SCENARIOS entry from its spine pack — every totals[*][4]
    and lines[*][3] figure is formatted from the engine, never typed."""
    spine = SPINES[key]
    inputs = spine["inputs"]
    bill = spine["bill"]
    contracted = inputs["contractedKwh"]
    total = inputs["totalConsumptionKwh"]
    shortfall = total - contracted
    fmp = inputs["fmp"]
    strike = inputs["strikePrice"]
    has_excess = "excess" in spine

    if has_excess:
        given_en = (f"Q_KH = {fmt_int(total)}   ·   gen ≈ {fmt_int(spine['excess']['generationKwh'])}"
                    f"   ·   FMP = {fmt_int(fmp)}   ·   Strike = {fmt_int(strike)}")
        given_vi = (f"Q_KH = {fmt_int(total)}   ·   phát ≈ {fmt_int(spine['excess']['generationKwh'])}"
                    f"   ·   FMP = {fmt_int(fmp)}   ·   Giá TH = {fmt_int(strike)}")
    else:
        given_en = (f"Q_c = {fmt_int(contracted)}   ·   Q_KH = {fmt_int(total)}"
                    f"   ·   FMP = {fmt_int(fmp)}   ·   Strike = {fmt_int(strike)}")
        given_vi = (f"Q_c = {fmt_int(contracted)}   ·   Q_KH = {fmt_int(total)}"
                    f"   ·   FMP = {fmt_int(fmp)}   ·   Giá TH = {fmt_int(strike)}")

    lines = [
        ("1", "Market energy", "Điện năng thị trường", f"{fmt_int(contracted)} × {fmt_int(fmp)} × 1.026 × 1.008"),
        ("2", "Service fee", "Phí dịch vụ", f"{fmt_int(contracted)} × {fmt_int(SERVICE)}"),
        ("3", "Clearing fee", "Phí bù trừ", f"{fmt_int(contracted)} × {CLEARING:.2f}"),
        ("4", "Additional retail", "Mua thêm bán lẻ", f"{fmt_int(shortfall)} × {fmt_int(RETAIL)}"),
    ]

    cfd_formula = f"({fmt_int(strike)} − {fmt_int(fmp)}) × {fmt_int(contracted)}"
    totals = [
        ("C_EVN", "C_EVN", "lines 1 + 2 + 3 + 4", "dòng 1 + 2 + 3 + 4", fmt_vnd(bill["cEvn"]["vnd"])),
        ("5", "CfD", cfd_formula, cfd_formula, fmt_signed_vnd(bill["lines"]["cfd"]["vnd"])),
        ("C_KH", "C_KH", "C_EVN + CfD", "C_EVN + CfD", fmt_vnd(bill["cKh"]["vnd"])),
        ("", "Effective", f"C_KH ÷ {fmt_int(total)}", f"C_KH ÷ {fmt_int(total)}", fmt_effective(bill["cKh"]["vnd"], total)),
    ]

    excess_block = None
    if has_excess:
        ex = spine["excess"]
        excess_block = [
            ("Excess volume", "Sản lượng dư thừa", f"{fmt_int(ex['generationKwh'])} − {fmt_int(total)}", ""),
            ("Spot value of excess", "Giá spot của dư thừa", ex["spotFormulaText"], ""),
            ("Foregone CfD uplift", "CfD bị bỏ lỡ", f"({fmt_int(strike)} − {fmt_int(fmp)}) × {fmt_int(ex['excessKwh'])}", ""),
        ]

    return {
        "key": key.upper(),
        "en_title": en_title,
        "vi_title": vi_title,
        "given_en": given_en,
        "given_vi": given_vi,
        "lines": lines,
        "totals": totals,
        "excess_block": excess_block,
    }


SCENARIOS = [
    build_scenario("s1", "SCENARIO 1 — MATCHED  (Load = Gen)", "KỊCH BẢN 1 — KHỚP  (Phụ tải = Phát điện)"),
    build_scenario("s2", "SCENARIO 2 — SHORTFALL  (Load > Gen)", "KỊCH BẢN 2 — THIẾU HỤT  (Phụ tải > Phát điện)"),
    build_scenario("s3", "SCENARIO 3 — EXCESS  (Load < Gen)", "KỊCH BẢN 3 — DƯ THỪA  (Phụ tải < Phát điện)"),
]

# 3-case comparison summary rows (header + 3 data rows; EN | VI)
COMPARISON_HEADER = [
    "Scenario / Kịch bản",
    "Volume axis\nTrục sản lượng",
    "Line 4 (retail)\nDòng 4 (bán lẻ)",
    "FMP vs strike\nFMP so với giá TH",
    "CfD\nCfD",
    "C_EVN\nC_EVN",
    "C_KH\nC_KH",
    "Effective (VND/kWh)\nHiệu dụng",
    "Risk lesson\nBài học rủi ro",
]
# Per-scenario narrative fields not carried by the spine (labels, volume-axis
# framing, cross-payer flow direction, the risk-lesson teaching line).
COMPARISON_META = {
    "s1": {
        "label": "S1 Matched  |  S1 Khớp",
        "volume_axis": "Q_c = Q_KH",
        "flow_en": "factory → developer",
        "flow_vi": "nhà máy → nhà phát triển",
        "risk_lesson": "strike must be bankable  |  giá thực hiện phải khả thi",
    },
    "s2": {
        "label": "S2 Shortfall  |  S2 Thiếu hụt",
        "volume_axis": "Q_c < Q_KH",
        "flow_en": "developer → factory",
        "flow_vi": "nhà phát triển → nhà máy",
        "risk_lesson": "shortfall always at retail  |  phần thiếu luôn tính giá bán lẻ",
    },
    "s3": {
        "label": "S3 Excess  |  S3 Dư thừa",
        "volume_axis": "gen > Q_KH (over-contract)",
        "flow_en": "factory → developer",
        "flow_vi": "nhà máy → nhà phát triển",
        "risk_lesson": "excess earns nothing  |  phần dư không tạo giá trị",
    },
}


def build_comparison_rows() -> list[list[str]]:
    rows = []
    for key in ("s1", "s2", "s3"):
        spine = SPINES[key]
        meta = COMPARISON_META[key]
        inputs = spine["inputs"]
        bill = spine["bill"]
        total = inputs["totalConsumptionKwh"]
        fmp = inputs["fmp"]
        strike = inputs["strikePrice"]
        cfd = bill["lines"]["cfd"]["vnd"]
        additional = bill["lines"]["additionalPurchase"]["vnd"]
        effective = round(bill["cKh"]["vnd"] / total)
        rows.append([
            meta["label"],
            meta["volume_axis"],
            fmt_int(additional),
            f"{fmt_int(fmp)} {'<' if fmp < strike else '>'} {fmt_int(strike)}",
            f"{fmt_signed_vnd(cfd)}\n{meta['flow_en']}\n{meta['flow_vi']}",
            fmt_int(bill["cEvn"]["vnd"]),
            fmt_int(bill["cKh"]["vnd"]),
            f"~{effective:,}",
            meta["risk_lesson"],
        ])
    return rows


COMPARISON_ROWS = build_comparison_rows()

# Negotiation grid — blank proposal rows + guidance text
NEGOTIATION_GRID_HEADER = [
    "Round  |  Lượt",
    "Proposed strike (VND/kWh)  |  Giá TH đề xuất",
    "CfD = (strike − 1,150) × 5,000,000  |  CfD",
    "Flows to…  |  Chảy về…",
]
NEGOTIATION_GRID_ROWS = [
    "Off-taker opens  |  Bên mua mở",
    "Developer opens  |  Nhà phát triển mở",
    "Counter  |  Phản hồi",
    "Agreed strike  |  Giá TH thống nhất",
]
NEGOTIATION_RESULT_ROWS = [
    "Resulting C_KH at the agreed strike  |  C_KH kết quả tại giá TH thống nhất",
    "Strike where the CfD crosses zero (= FMP)  |  Giá TH nơi CfD cắt qua không (= FMP)",
    "Strike where the buyer first beats BAU (app multi-year)  |  Giá TH nơi bên mua lần đầu thắng BAU (bảng nhiều năm)",
]
NEGOTIATION_GUIDANCE_EN = (
    "Worked example: at strike = 1,200 VND/kWh and FMP = 1,150, CfD = (1,200 − 1,150) × 5,000,000 = "
    "+250,000,000 VND (factory tops up developer). The CfD crosses zero at strike = FMP. The three gates — "
    "buyer ≤ BAU, seller IRR, lender DSCR — are usually only all green inside a narrow strike band. "
    "Drag the strike slider in Workshop 1 to test your agreed value."
)
NEGOTIATION_GUIDANCE_VI = (
    "Ví dụ minh hoạ: ở giá TH = 1.200 VND/kWh và FMP = 1.150, CfD = (1.200 − 1.150) × 5.000.000 = "
    "+250.000.000 VND (nhà máy bù cho nhà phát triển). CfD cắt qua không khi giá TH = FMP. Ba cổng — "
    "mua ≤ BAU, bán IRR, ngân hàng DSCR — thường chỉ cùng xanh trong một dải giá TH hẹp. "
    "Kéo thanh trượt giá thực hiện trong Workshop 1 để kiểm tra giá trị đã thoả thuận."
)


# ----------------------------------------------------------------------
# XML helpers
# ----------------------------------------------------------------------
def _w(tag: str) -> str:
    return qn(f"w:{tag}")


def _make(tag: str, **attrs) -> OxmlElement:
    el = OxmlElement(f"w:{tag}")
    for k, v in attrs.items():
        if v is None:
            continue
        el.set(_w(k), str(v))
    return el


def set_cell_shading(cell, fill_hex: str) -> None:
    """Set the background fill on a ``w:tc`` (table cell)."""
    tcPr = cell._tc.get_or_add_tcPr()
    # remove any prior shd
    for old in tcPr.findall(_w("shd")):
        tcPr.remove(old)
    shd = _make("shd", val="clear", color="auto", fill=fill_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, left=150, bottom=80, right=150) -> None:
    """Set cell padding (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(_w("tcMar")):
        tcPr.remove(old)
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        m = _make(side, w=str(val), type="dxa")
        mar.append(m)
    tcPr.append(mar)


def set_cell_borders(cell, color="cccccc", sz=4) -> None:
    """Add a thin gray border to a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(_w("tcBorders")):
        tcPr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = _make(side, val="single", sz=str(sz), space="0", color=color)
        borders.append(b)
    tcPr.append(borders)


def set_cell_width(cell, dxa: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(_w("tcW")):
        tcPr.remove(old)
    w = _make("tcW", w=str(dxa), type="dxa")
    tcPr.append(w)


def set_table_borders(table, color="000000", sz=4) -> None:
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(_w("tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_make(side, val="single", sz=str(sz), space="0", color=color))
    tblPr.append(borders)


def set_table_width(table, dxa: int) -> None:
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(_w("tblW")):
        tblPr.remove(old)
    tblPr.append(_make("tblW", w=str(dxa), type="dxa"))
    # also fix grid
    grid = table._tbl.find(_w("tblGrid"))
    if grid is not None:
        for old in list(grid):
            grid.remove(old)
        grid.append(_make("gridCol", w=str(dxa)))


def set_para_shading(para, fill_hex: str) -> None:
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(_w("shd")):
        pPr.remove(old)
    pPr.append(_make("shd", val="clear", color="auto", fill=fill_hex))


def set_para_alignment(para, val: str) -> None:
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(_w("jc")):
        pPr.remove(old)
    pPr.append(_make("jc", val=val))


def add_run(para, text: str, *, bold: bool = False, color: str | None = None,
            size_half_pt: str | None = None, font: str | None = None) -> None:
    run = para.add_run(text)
    rPr = run._r.get_or_add_rPr()
    if bold:
        b = OxmlElement("w:b")
        b.set(_w("val"), "1")
        rPr.append(b)
        bcs = OxmlElement("w:bCs")
        bcs.set(_w("val"), "1")
        rPr.append(bcs)
    if color:
        c = _make("color", val=color)
        rPr.append(c)
    if size_half_pt:
        rPr.append(_make("sz", val=size_half_pt))
        rPr.append(_make("szCs", val=size_half_pt))
    if font:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(_w("ascii"), font)
        rFonts.set(_w("hAnsi"), font)
        rPr.append(rFonts)
    # disable RTL fallback noise
    rPr.append(_make("rtl", val="0"))


def add_page_break(para) -> None:
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(_w("type"), "page")
    run._r.append(br)


# ----------------------------------------------------------------------
# Builder primitives
# ----------------------------------------------------------------------
def banner(doc, en_title: str, vi_title: str, legal: str | None = None) -> None:
    """Three-paragraph centered banner matching the reference: white 18pt
    bold title + D6E4F0 14pt subtitle (+ optional 9pt legal line)."""
    p1 = doc.add_paragraph()
    set_para_shading(p1, BANNER_FILL)
    set_para_alignment(p1, "center")
    add_run(p1, en_title, bold=True, color="FFFFFF", size_half_pt=SZ_TITLE)

    p2 = doc.add_paragraph()
    set_para_shading(p2, BANNER_FILL)
    set_para_alignment(p2, "center")
    add_run(p2, vi_title, bold=True, color=SUBTITLE_FILL, size_half_pt=SZ_SUBTITLE)

    if legal:
        p3 = doc.add_paragraph()
        set_para_shading(p3, BANNER_FILL)
        set_para_alignment(p3, "center")
        add_run(p3, legal, color=SUBTITLE_FILL, size_half_pt=SZ_LEGAL)

    # spacer
    doc.add_paragraph()


def section_head(doc, text: str) -> None:
    p = doc.add_paragraph()
    add_run(p, text, bold=True, color=SECTION_COLOR, size_half_pt=SZ_SECTION)


def subhead(doc, text: str) -> None:
    p = doc.add_paragraph()
    add_run(p, text, bold=True, color=SUBHEAD_COLOR, size_half_pt=SZ_BODY)


def body_para(doc, text: str, *, italic: bool = False, color: str = "000000") -> None:
    p = doc.add_paragraph()
    add_run(p, text, color=color, size_half_pt=SZ_BODY)
    if italic:
        p.runs[-1].italic = True


def data_table(doc, rows: list[list[str]], *, header_fill: str = HEADER_FILL,
               col_widths: list[int] | None = None, has_header: bool = True) -> None:
    """Build a bordered table; first row = white-on-blue header when has_header."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    n_rows = len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = False
    set_table_borders(table)
    if col_widths is None:
        col_widths = [int(TABLE_WIDTH) // n_cols] * n_cols
    set_table_width(table, sum(col_widths))
    grid = table._tbl.find(_w("tblGrid"))
    for old in list(grid):
        grid.remove(old)
    for w in col_widths:
        grid.append(_make("gridCol", w=str(w)))
    for ri, row in enumerate(rows):
        for ci, txt in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            for line in str(txt).split("\n"):
                if p.text:
                    p = cell.add_paragraph()
                add_run(p, line, color="000000", size_half_pt=SZ_BODY)
            set_cell_width(cell, col_widths[ci])
            set_cell_margins(cell)
            set_cell_borders(cell)
            if has_header and ri == 0:
                set_cell_shading(cell, header_fill)
                # recolor all runs in this header cell to white + bold
                for p2 in cell.paragraphs:
                    for run in p2.runs:
                        rPr = run._r.get_or_add_rPr()
                        # remove existing color
                        for old in rPr.findall(_w("color")):
                            rPr.remove(old)
                        rPr.append(_make("color", val="FFFFFF"))
                        b = OxmlElement("w:b")
                        b.set(_w("val"), "1")
                        rPr.append(b)
            else:
                set_cell_shading(cell, "FFFFFF")


def callout(doc, en_text: str, vi_text: str, *, fill: str = SUBTITLE_FILL) -> None:
    """A 1×2 EN|VI callout box (light-blue background)."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_borders(table, color="cccccc", sz=4)
    half = int(TABLE_WIDTH) // 2
    set_table_width(table, int(TABLE_WIDTH))
    grid = table._tbl.find(_w("tblGrid"))
    for old in list(grid):
        grid.remove(old)
    grid.append(_make("gridCol", w=str(half)))
    grid.append(_make("gridCol", w=str(half)))
    en_cell, vi_cell = table.cell(0, 0), table.cell(0, 1)
    for cell, txt in ((en_cell, en_text), (vi_cell, vi_text)):
        cell.text = ""
        p = cell.paragraphs[0]
        add_run(p, txt, color="000000", size_half_pt=SZ_BODY)
        set_cell_width(cell, half)
        set_cell_margins(cell, top=120, left=200, bottom=120, right=200)
        set_cell_borders(cell, color="cccccc", sz=4)
        set_cell_shading(cell, fill)
    doc.add_paragraph()


def worksheet_grid(doc, scenario: dict) -> None:
    """A blank fillable 5-line compute grid (line / EN|VI label / calculation
    formula) plus an empty shaded answer column. Mirrors lessons/0011."""
    s = scenario
    rows = [["#", "Line  |  Dòng", "Calculation  |  Phép tính", "VND / month  |  VND / tháng"]]
    for num, en, vi, calc in s["lines"]:
        rows.append([num, f"{en}  |  {vi}", calc, ""])
    # C_EVN / CfD / C_KH / Effective blank rows
    for num, en, calc_en, calc_vi, _ans in s["totals"]:
        rows.append([num, f"**{en}**  |  **{en}**" if en.startswith("C_") else en, f"{calc_en}", ""])

    table = doc.add_table(rows=len(rows), cols=4)
    table.autofit = False
    set_table_borders(table)
    col_widths = [int(COL_LABEL), int(COL_LABEL) * 2 + 400, int(COL_CALC), int(COL_ANSWER)]
    set_table_width(table, sum(col_widths))
    grid = table._tbl.find(_w("tblGrid"))
    for old in list(grid):
        grid.remove(old)
    for w in col_widths:
        grid.append(_make("gridCol", w=str(w)))
    for ri, row in enumerate(rows):
        for ci, txt in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            for line in str(txt).split("\n"):
                if p.text or line != str(txt).split("\n")[0]:
                    p = cell.add_paragraph()
                add_run(p, line, color="000000", size_half_pt=SZ_BODY)
            set_cell_width(cell, col_widths[ci])
            set_cell_margins(cell)
            set_cell_borders(cell, color="cccccc", sz=4)
            if ri == 0:
                # header
                set_cell_shading(cell, HEADER_FILL)
                for p2 in cell.paragraphs:
                    for run in p2.runs:
                        rPr = run._r.get_or_add_rPr()
                        for old in rPr.findall(_w("color")):
                            rPr.remove(old)
                        rPr.append(_make("color", val="FFFFFF"))
                        b = OxmlElement("w:b")
                        b.set(_w("val"), "1")
                        rPr.append(b)
            else:
                # blank answer cells (col 3) get a fillable tint
                if ci == 3:
                    set_cell_shading(cell, "FFFEF2")
                else:
                    set_cell_shading(cell, "FFFFFF")
                # bold the C_EVN / C_KH label cells
                if ci == 1 and rows[ri][0] in ("C_EVN", "C_KH"):
                    for p2 in cell.paragraphs:
                        for run in p2.runs:
                            rPr = run._r.get_or_add_rPr()
                            b = OxmlElement("w:b")
                            b.set(_w("val"), "1")
                            rPr.append(b)
                    if ci == 3:
                        set_cell_shading(cell, "FFF7E6")
    doc.add_paragraph()


def totals_table(doc, scenario: dict) -> None:
    """A small 3-row table with C_EVN / CfD / C_KH + effective."""
    s = scenario
    rows = [
        ["Item  |  Mục", "Value (VND)  |  Giá trị"],
        [f"C_EVN", s["totals"][0][4]],
        [f"CfD  |  CfD", s["totals"][1][4]],
        [f"C_KH", s["totals"][2][4]],
        ["Effective (VND/kWh)  |  Hiệu dụng", s["totals"][3][4]],
    ]
    col_widths = [int(TABLE_WIDTH) // 2, int(TABLE_WIDTH) - int(TABLE_WIDTH) // 2]
    data_table(doc, rows, col_widths=col_widths)


def excess_block(doc, scenario: dict) -> None:
    """S3 only — blank fillable excess analysis (volume / spot value / foregone CfD)."""
    rows = [["Item  |  Mục", "Formula  |  Công thức", "Value (VND)  |  Giá trị"]]
    for en, vi, calc, _ in scenario["excess_block"]:
        rows.append([f"{en}  |  {vi}", calc, ""])
    col_widths = [int(COL_LABEL) * 2 + 400, int(COL_CALC), int(COL_ANSWER)]
    data_table(doc, rows, col_widths=col_widths)


def footer_line(doc, text: str) -> None:
    p = doc.add_paragraph()
    set_para_alignment(p, "center")
    add_run(p, text, color=FOOTER_COLOR, size_half_pt=SZ_FOOTER)


# ----------------------------------------------------------------------
# Page-level emitters
# ----------------------------------------------------------------------
def emit_title_banner(doc) -> None:
    banner(
        doc,
        en_title="DPPA Worksheets & Answers",
        vi_title="Phiếu tính DPPA & Đáp án",
        legal=(
            "Legal basis: Decree 57/2025/NĐ-CP (amended by Decree 243/2026/NĐ-CP)  |  "
            "Căn cứ pháp lý: Nghị định 57/2025/NĐ-CP (sửa đổi bởi Nghị định 243/2026/NĐ-CP)"
        ),
    )


def emit_constants_block(doc) -> None:
    section_head(doc, "SHARED CONSTANTS  |  HẰNG SỐ CHUNG")
    rows = [
        ["Symbol  |  Ký hiệu", "Value  |  Giá trị", "Meaning  |  Ý nghĩa"],
        ["k × K_pp", f"{K_KPP}", "combined loss factor (Decree 57/2025)  |  hệ số tổn thất kết hợp"],
        ["C_dppa_dv (service)", f"{SERVICE} VND/kWh", "power-system service fee (line 2)  |  phí dịch vụ hệ thống"],
        ["P_cl (clearing)", f"{CLEARING} VND/kWh", "differential clearing fee (line 3)  |  phí bù trừ chênh lệch"],
        ["fees (service + clearing)", f"{FEES} VND/kWh", "fixed DPPA fees on every matched kWh  |  phí cố định trên mỗi kWh khớp"],
        ["P1 (retail)", f"{RETAIL:,} VND/kWh", "residual retail price (line 4)  |  giá bán lẻ phần thiếu hụt"],
    ]
    col_widths = [int(TABLE_WIDTH) // 4, int(TABLE_WIDTH) // 4,
                  int(TABLE_WIDTH) - int(TABLE_WIDTH) // 4 * 2]
    data_table(doc, rows, col_widths=col_widths)
    body_para(doc, "All totals below are reproduced by buildFiveLineBill() in app/src/modules/settlement.js (k × K_pp = 1.026 × 1.008 = 1.034208).",
              italic=True, color="555555")


def emit_scenario(doc, scenario: dict) -> None:
    section_head(doc, scenario["en_title"] + "   |   " + scenario["vi_title"])

    # Inputs line (callout)
    callout(doc,
            en_text="Inputs: " + scenario["given_en"],
            vi_text="Thông số: " + scenario["given_vi"])

    subhead(doc, "Worksheet  |  Phiếu tính")
    worksheet_grid(doc, scenario)

    subhead(doc, "Answer (totals)  |  Đáp án")
    totals_table(doc, scenario)

    if scenario["excess_block"]:
        subhead(doc, "The excess (settles nothing)  |  Phần dư thừa (không tạo giá trị)")
        excess_block(doc, scenario)

    # page break
    p = doc.add_paragraph()
    add_page_break(p)


def emit_negotiation(doc) -> None:
    banner(doc,
           en_title="NEGOTIATION  |  THE STRIKE",
           vi_title="ĐÀM PHÁN  |  GIÁ THỰC HIỆN",
           legal="Starting point: Scenario 1 volumes (5,000,000 kWh, FMP ≈ 1,150)  |  Điểm khởi đầu: sản lượng Kịch bản 1")

    section_head(doc, "Proposal grid  |  Bảng đề xuất")
    rows = [NEGOTIATION_GRID_HEADER]
    for r in NEGOTIATION_GRID_ROWS:
        rows.append([r, "", "", ""])
    data_table(doc, rows)

    section_head(doc, "Resulting economics  |  Kết quả kinh tế")
    rows = [["Metric  |  Chỉ tiêu", "Value  |  Giá trị"]]
    for r in NEGOTIATION_RESULT_ROWS:
        rows.append([r, ""])
    col_widths = [int(TABLE_WIDTH) * 3 // 5, int(TABLE_WIDTH) - int(TABLE_WIDTH) * 3 // 5]
    data_table(doc, rows, col_widths=col_widths)

    section_head(doc, "Guidance  |  Hướng dẫn")
    callout(doc, NEGOTIATION_GUIDANCE_EN, NEGOTIATION_GUIDANCE_VI)

    p = doc.add_paragraph()
    add_page_break(p)


def emit_comparison(doc) -> None:
    section_head(doc, "SCENARIO COMPARISON SUMMARY  |  BẢNG SO SÁNH CÁC KỊCH BẢN")
    rows = [COMPARISON_HEADER] + COMPARISON_ROWS
    data_table(doc, rows)


# ----------------------------------------------------------------------
# Main build
# ----------------------------------------------------------------------
def clear_body_preserve_sectpr(doc) -> None:
    body = doc.element.body
    sectPr = body.find(_w("sectPr"))
    for child in list(body):
        if child.tag == _w("sectPr"):
            continue
        body.remove(child)
    # ensure sectPr is still last
    if sectPr is not None and body[-1].tag != _w("sectPr"):
        body.append(sectPr)


def build() -> None:
    if not TEMPLATE.exists():
        raise SystemExit(f"Template not found: {TEMPLATE}")
    shutil.copy(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))
    clear_body_preserve_sectpr(doc)

    emit_title_banner(doc)
    emit_constants_block(doc)

    p = doc.add_paragraph()
    add_page_break(p)

    for s in SCENARIOS:
        emit_scenario(doc, s)

    emit_negotiation(doc)
    emit_comparison(doc)

    footer_line(doc, FOOTER_TEXT)
    doc.save(str(OUTPUT))
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    build()
